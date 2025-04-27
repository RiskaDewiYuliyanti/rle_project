from docx import Document
import re

def rle_compress(text):
    if not text:
        return ""
    compressed = []
    count = 1
    prev_char = text[0]

    for char in text[1:]:
        if char == prev_char:
            count += 1
        else:
            compressed.append(f"{count}{prev_char}")
            count = 1
            prev_char = char
    compressed.append(f"{count}{prev_char}")
    return ''.join(compressed)

def rle_decompress(text):
    decompressed = []
    tokens = re.findall(r'(\d+)(\D)', text)
    for count, char in tokens:
        decompressed.append(int(count) * char)
    return ''.join(decompressed)

def read_docx(file_path):
    doc = Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def save_to_docx(text, output_filename):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_filename)